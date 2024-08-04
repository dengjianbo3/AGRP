from typing import List, Optional
from langchain.docstore.document import Document
from langchain.document_loaders.unstructured import UnstructuredFileLoader
from tqdm import tqdm
import docx
import subprocess
from .splitter import ChineseRecursiveTextSplitter
import re

class DocxDocLoader(UnstructuredFileLoader):
    def __init__(
        self,
        file_path: str,
        file_name: str = None, 
        chunk_size: int = 50,
        chunk_overlap: int = 0,
        unstructured_kwargs: Optional[dict] = None,
        dataclean: bool = True,  # 添加数据清理开关
        add_file_name: bool = True,
    ):
        super().__init__(
            file_path=file_path,
            unstructured_kwargs=unstructured_kwargs,
        )
        self.text_splitter = ChineseRecursiveTextSplitter(keep_separator=True, is_separator_regex=True, chunk_size=chunk_size, chunk_overlap=chunk_overlap)
        self.dataclean = dataclean  
        self.add_file_name = add_file_name
        self.file_name = file_name

    def load(self) -> List[Document]:
        """Load data into document objects."""
        docs = []
        try:
            text = self.docx2text(self.file_path)
            if self.dataclean:
                text = self.clean_data(text)
            metadata = {"source": self.file_path}
            chunks = self.text_splitter.split_text(text)
            for chunk in chunks:
                if self.add_file_name:
                    chunk = f"[{self.file_name}] : {chunk}"
                doc = Document(page_content=chunk, metadata=metadata)
                docs.append(doc)
            return docs
        except Exception as e:
            raise RuntimeError(f"Error loading {self.file_path}") from e

    def docx2text(self, filepath):
        doc = docx.Document(filepath)
        fullText = []
        for para in doc.paragraphs:
            if not para._p.xpath('.//w:drawing'):  # 检查段落中是否包含图片
                fullText.append(para.text)
        return "\n".join(fullText)

    def doc2text(self, filepath):
        output = subprocess.check_output(['antiword', filepath])
        return output.decode('utf-8')

    def clean_data(self, text: str) -> str:  # 更新数据清理方法
        # 去除特殊符号
        # text = re.sub(r'[^\w\s【】：，。；！：]', '', text)
        # 替换连续多个空格为单个空格
        text = re.sub(r'\s+', ' ', text)
        # 移除多余的空行
        lines = text.split("\n")
        non_empty_lines = [line.strip() for line in lines if line.strip()]
        # 清理特定无意义的格式或标识
        cleaned_lines = []
        for line in non_empty_lines:
            if "Image Format" not in line and "Dimensions" not in line and "ColorSpace" not in line and "ExifIFD" not in line:
                cleaned_lines.append(line)
        return "\n".join(cleaned_lines)


class RapidOCRLoader(UnstructuredFileLoader):
    def _get_elements(self) -> List:
        def img2text(filepath):
            from rapidocr_onnxruntime import RapidOCR
            resp = ""
            ocr = RapidOCR()
            result, _ = ocr(filepath)
            if result:
                ocr_result = [line[1] for line in result]
                resp += "\n".join(ocr_result)
            return resp

        text = img2text(self.file_path)
        from unstructured.partition.text import partition_text
        return partition_text(text=text, **self.unstructured_kwargs)
    

class RapidOCRPDFLoader(UnstructuredFileLoader):
    def __init__(
        self,
        file_path: str,
        file_name: str = None,
        chunk_size: int = 50,
        chunk_overlap: int = 0,
        unstructured_kwargs: Optional[dict] = None,
        dataclean: bool = True,
        add_file_name: bool = True
    ):
        super().__init__(
            file_path=file_path,
            unstructured_kwargs=unstructured_kwargs,
        )
        self.text_splitter = ChineseRecursiveTextSplitter(keep_separator=True, is_separator_regex=True, chunk_size=chunk_size, chunk_overlap=chunk_overlap)

    def load(self) -> List[Document]:
        """Load data into document objects."""
        docs = []
        try:
            text = self.pdf2text(self.file_path)
            metadata = {"source": self.file_path}
            chunks = self.text_splitter.split_text(text)
            for chunk in chunks:
                doc = Document(page_content=chunk, metadata=metadata)
                docs.append(doc)
        except Exception as e:
            raise RuntimeError(f"Error loading {self.file_path}") from e
        return docs

    def pdf2text(self, filepath):
        import fitz  # pyMuPDF里面的fitz包，不要与pip install fitz混淆
        from rapidocr_onnxruntime import RapidOCR
        import numpy as np
        ocr = RapidOCR()
        doc = fitz.open(filepath)
        resp = ""
        b_unit = tqdm(total=doc.page_count, desc="RapidOCRPDFLoader context page index: 0")
        for i, page in enumerate(doc):
            # 更新描述
            b_unit.set_description("RapidOCRPDFLoader context page index: {}".format(i))
            # 立即显示进度条更新结果
            b_unit.refresh()
            # TODO: 依据文本与图片顺序调整处理方式
            text = page.get_text("")
            resp += text + "\n"
            img_list = page.get_images()
            for img in img_list:
                pix = fitz.Pixmap(doc, img[0])
                img_array = np.frombuffer(pix.samples, dtype=np.uint8).reshape(pix.height, pix.width, -1)
                result, _ = ocr(img_array)
                if result:
                    ocr_result = [line[1] for line in result]
                    resp += "\n".join(ocr_result)
            # 更新进度
            b_unit.update(1)
        return resp